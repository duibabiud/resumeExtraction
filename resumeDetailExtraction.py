
import os, glob, re
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import StringIO
import pandas as pd
from docx import Document



def convertDocxToText(path):
    txt = []
    document = Document(path)
    txt.append("\n".join([para.text for para in document.paragraphs]))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                txt.append("\n".join([para.text for para in cell.paragraphs]))

    return ('\n'.join([str(e) for e in txt]))



def convertPDFToText(path):
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    codec = 'utf-8'
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, laparams=laparams)
    fp = open(path, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password = ""
    maxpages = 0
    caching = True
    pagenos = set()
    for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password, caching=caching,
                                  check_extractable=True):
        interpreter.process_page(page)
    fp.close()
    device.close()
    string = retstr.getvalue()
    retstr.close()
    return string


class exportToCSV:
    def __init__(self, fileName='resultsCSV.csv', resetFile=False):
        headers = [
            'EMPLOYEE NAME',
            'DATE OF BIRTH',
            'GENDER',
            'NATIONALITY',
            'CURRENT ADDRESS'
        ]
        if not os.path.isfile(fileName) or resetFile:
            # Will create/reset the file as per the evaluation of above condition
            fOut = open(fileName, 'w')
            fOut.close()
        fIn = open(fileName)  ########### Open file if file already present
        inString = fIn.read()
        fIn.close()
        if len(inString) <= 0:  ######### If File already exsists but is empty, it adds the header
            fOut = open(fileName, 'w')
            fOut.write(','.join(headers) + '\n')
            fOut.close()

    def write(self, infoDict):
        fOut = open('resultsCSV.csv', 'a+')
        # Individual elements are dictionaries
        writeString = ''
        try:
            writeString += str(infoDict['EMPLOYEE NAME']) + ','
            writeString += str(infoDict['DATE OF BIRTH']) + ","
            writeString += str(infoDict['GENDER']) + ","
            writeString += str(infoDict['NATIONALITY']) + ","
            writeString += str(infoDict['CURRENT ADDRESS']) + ",\n"
            fOut.write(writeString)
        except:
            fOut.write('FAILED_TO_WRITE\n')
        fOut.close()


class Parse():
    information = []
    inputString = ''

    def __init__(self):
        print('Starting Program...')

        # Glob module matches certain patterns
        doc_files = glob.glob("resumes/*.doc")
        docx_files = glob.glob("resumes/*.docx")
        pdf_files = glob.glob("resumes/*.pdf")

        files = set(doc_files + docx_files + pdf_files)
        files = list(files)
        print("%d files identified" % len(files))

        for f in files:
            print("Reading File %s" % f)
            # info is a dictionary that stores all the data obtained from parsing
            info = {}

            self.inputString = self.readFile(f)
            info['Filename'] = f

            self.getName(self.inputString, info)
            self.getDOB(self.inputString, info)
            self.getGender(self.inputString, info)
            self.getNationality(self.inputString, info)
            self.getCurrentAddress(self.inputString, info)

            csv = exportToCSV()
            csv.write(info)
            self.information.append(info)
            print(info)

    def readFile(self, fileName):

        extension = fileName.split(".")[-1]
        if extension == "pdf":
            return convertPDFToText(fileName)
        elif extension == "docx":
            try:
                return convertDocxToText(fileName)
            except:
                return ''
                pass
        else:
            print('Unsupported format')
            return ''

    #Function to get Employee Name.
    def getName(self, inputString, infoDict):
        #nm = re.search(r'([a-zA-Z0-9]*@[a-z]*.((com)|(org)|(net)|(co.in)))',inputString,re.IGNORECASE) # incomplete
        #eml = re.search(r'\S+@\S+', inputString)
        #if nm is None:
        infoDict['EMPLOYEE NAME'] = 'NAME'
        #else:
        #infoDict['EMPLOYEE NAME'] = nm[0].split('@')[0]
        #else:
        #   infoDict['EMPLOYEE NAME'] = nm

    #Function to get Date Of birth.
    def getDOB(self, inputString, infoDict):
        dob = re.search(r'(((DOB\s*:)|(D.O.B\s*:)|(date of birth\s*:))[- ]\s*\d{1,2}[-,./ ]*\s*((\d{1,2})|([a-zA-z]*))[-,./ ]\s*\d{4})|((DOB\s*:)|(D.O.B\s*:)|(date of birth\s*:))-*\s*(\d{1,2}((nd)|(th)|(rd))[-,./ ]*\s*[a-zA-z]*[-,./ ]*\s*\d{4})|((DOB)|(D.O.B)|(Date of Birth))\s*:-*\s*\d{1,2}[-,./ ]\s*\d{1,2}[-,./ ]\s*\d{4}', inputString, re.IGNORECASE)
        #dob = re.search(r'\d{1,2}[-,./ ]\s*((\d{1,2})|([a-zA-z]*))[-,./ ]\s*\d{4}|\d{1,2}[a-z]*[-,./ ]\s*[a-z]*[-,./ ]\s*\d{4}|[a-zA-Z]+[-,./ ]\s*\d{1,2}[-,./ ]\s*\d{4}', inputString, re.IGNORECASE)
        if dob is None:
            infoDict['DATE OF BIRTH'] = 'NA'
        else:
            d = dob[0].split(':')
            if len(d) == 2:
                d = d[1].strip()
                infoDict['DATE OF BIRTH'] = d.replace(',','')
            else:
                d = dob[0].strip()
                infoDict['DATE OF BIRTH'] = d.replace(',','')

    #Function to get Gender.
    def getGender(self, inputString, infoDict):
        g = re.search(r'((((gender)|(sex))[\s]*:\s*((male)|(female)|(m)|(f)))|((male)|(female)))', inputString, re.IGNORECASE)
        if g is None:
            infoDict['GENDER'] = 'NA'
        else:
            g = g[0].split(':')
            if len(g) == 2:
                infoDict['GENDER'] = g[1].strip()
            else:
                infoDict['GENDER'] = g[0]

    #Function to get Nationality.
    def getNationality(self, inputString, infoDict):
        nl = re.search(r'((nationality)|(country))[\s]*:\s*[a-zA-Z]*', inputString, re.IGNORECASE)
        if nl is None:
            infoDict['NATIONALITY'] = 'NA'
        else:
            nl = nl[0].split(':')
            if len(nl) == 2:
                infoDict['NATIONALITY'] = nl[1].strip()
            else:
                infoDict['NATIONALITY'] = nl[0]

    #Function to get Current Address
    def getCurrentAddress(self, inputString, infoDict):
        result = re.search(r'(((place)|(current location)|(current address))\s*:\s[A-Za-z]*)', inputString,re.IGNORECASE)

        if result is None:
            infoDict['CURRENT ADDRESS'] = 'NA'
        else:
            result = result[0].split(':')
            if len(result) == 2:
                infoDict['CURRENT ADDRESS'] = result[1].strip()
            else:
                infoDict['CURRENT ADDRESS'] = result[0]

if __name__ == "__main__":
    p = Parse()
