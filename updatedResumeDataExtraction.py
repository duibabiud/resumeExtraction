import os, glob, re
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import StringIO
from docx import Document
import PyPDF2


def convertDocxToText(path):
    txt = []
    document = Document(path)
    txt.append("\n".join([para.text for para in document.paragraphs]))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                txt.append("\n".join([para.text for para in cell.paragraphs]))

    return ('\n'.join([str(e) for e in txt]))



def convertPDFToText(path):
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, laparams=laparams)
    fp = open(path, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password = ""
    maxpages = 0
    caching = True
    pagenos = set()
    for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password, caching=caching,
                                  check_extractable=True):
        interpreter.process_page(page)
    fp.close()
    device.close()
    string = retstr.getvalue()
    retstr.close()
    string = string.replace('\n','')
    return string

def convertPDFToTextUsingPypdf2(path):
    pdfFileObj = open(path, 'rb')
    pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
    num_pages = pdfReader.numPages
    count = 0
    text = ""

    while count < num_pages:
        pageObj = pdfReader.getPage(count)
        count += 1
        text += pageObj.extractText()
    return text

class exportToCSV:
    def __init__(self, fileName='resultsCSV.csv', resetFile=False):
        headers = [
            'EMPLOYEE NAME',
            'DATE OF BIRTH',
            'GENDER',
            'NATIONALITY',
            'CURRENT ADDRESS'
        ]
        if not os.path.isfile(fileName) or resetFile:
            # Will create/reset the file as per the evaluation of above condition
            fOut = open(fileName, 'w')
            fOut.close()
        fIn = open(fileName)  ########### Open file if file already present
        inString = fIn.read()
        fIn.close()
        if len(inString) <= 0:  ######### If File already exsists but is empty, it adds the header
            fOut = open(fileName, 'w')
            fOut.write(','.join(headers) + '\n')
            fOut.close()

    def write(self, infoDict):
        fOut = open('resultsCSV.csv', 'a+')
        # Individual elements are dictionaries
        writeString = ''
        try:
            writeString += str(infoDict['EMPLOYEE NAME']) + ','
            writeString += str(infoDict['DATE OF BIRTH']) + ","
            writeString += str(infoDict['GENDER']) + ","
            writeString += str(infoDict['NATIONALITY']) + ","
            writeString += str(infoDict['CURRENT ADDRESS']) + ",\n"
            fOut.write(writeString)
        except:
            fOut.write('FAILED_TO_WRITE\n')
        fOut.close()


class Parse():
    information = []
    inputString = ''

    def __init__(self):
        print('Starting Program...')

        # Glob module matches certain patterns
        doc_files = glob.glob("resumes/*.doc")
        docx_files = glob.glob("resumes/*.docx")
        pdf_files = glob.glob("resumes/*.pdf")

        files = set(doc_files + docx_files + pdf_files)
        files = list(files)
        print("%d files identified" % len(files))

        for file in files:
            print("Reading File %s" % file)
            # info is a dictionary that stores all the data obtained from parsing
            info = {}

            self.inputString = self.readFile(file)
            info['Filename'] = file

            self.getName(self.inputString, info, file)
            self.getDOB(self.inputString, info, file)
            self.getGender(self.inputString, info, file)
            self.getNationality(self.inputString, info, file)
            self.getCurrentAddress(self.inputString, info, file)

            csv = exportToCSV()
            csv.write(info)
            self.information.append(info)
            print(info)

    def readFile(self, fileName):

        extension = fileName.split(".")[-1]
        if extension == "pdf":
            return convertPDFToText(fileName)
        elif extension == "docx":
            try:
                return convertDocxToText(fileName)
            except:
                return ''
                pass
        else:
            print('Unsupported format')
            return ''

    #Function to get Employee Name.
    def getName(self, inputString, infoDict, file):  # incomplete function
        infoDict['EMPLOYEE NAME'] = 'NAME'


    #Function to get Date Of birth.
    def getDOB(self, inputString, infoDict, file):
        reg = re.search(r'(((DOB)|(D.O.B)|(date of birth)|(birthday))\s*:*[- ]\s*\d{1,2}[-,./ ]*\s*((\d{1,2})|([a-zA-z]*))[-,./ ]\s*\d{4})|((DOB)|(D.O.B)|(date of birth)|(birthday))\s*:*-*\s*(\d{1,2}((st)|(nd)|(th)|(rd))[-,./ ]*\s*[a-zA-z]*[-,./ ]*\s*\d{4})|((DOB)|(D.O.B)|(Date of Birth)|(birthday))\s*:*-*\s*\d{1,2}[-,./ ]\s*\d{1,2}[-,./ ]\s*\d{4}|((DOB)|(D.O.B)|(Date of Birth)|(birthday))\s*:*-*\s*[a-zA-Z]*[-,./ ]\d{1,2}[-,./ ]\s*\d{4}', inputString, re.IGNORECASE)

        ext = file[-4:]
        if reg is None and ext == '.pdf':
            text = convertPDFToTextUsingPypdf2(file)  # checking for text using PyPDF2 if not found using pdfminer
            text = text.replace('\n', '')
            reg = re.search(r'(((DOB)|(D.O.B)|(date of birth)|(birthday))\s*:*[- ]\s*\d{1,2}[-,./ ]*\s*((\d{1,2})|([a-zA-z]*))[-,./ ]\s*\d{4})|((DOB)|(D.O.B)|(date of birth)|(birthday))\s*:*-*\s*(\d{1,2}((st)|(nd)|(th)|(rd))[-,./ ]*\s*[a-zA-z]*[-,./ ]*\s*\d{4})|((DOB)|(D.O.B)|(Date of Birth)|(birthday))\s*:*-*\s*\d{1,2}[-,./ ]\s*\d{1,2}[-,./ ]\s*\d{4}|((DOB)|(D.O.B)|(Date of Birth)|(birthday))\s*:*-*\s*[a-zA-Z]*[-,./ ]\d{1,2}[-,./ ]\s*\d{4}',text, re.IGNORECASE)

            if reg is None:
                infoDict['DATE OF BIRTH'] = 'NA'
            else:
                dobReg = re.search(
                    r'(\d{1,2}[-,./ ]*\s*((\d{1,2})|([a-zA-z]*))[-,./ ]\s*\d{4})|(\d{1,2}((st)|(nd)|(th)|(rd))[-,./ ]*\s*[a-zA-z]*[-,./ ]*\s*\d{4})|(d{1,2}[-,./ ]\s*\d{1,2}[-,./ ]\s*\d{4}|[a-zA-Z]*[-,./ ]\d{1,2}[-,./ ]\s*\d{4})',
                    reg[0], re.IGNORECASE)
                infoDict['DATE OF BIRTH'] = dobReg[0].replace(',', ' ')

        elif reg is None:
            infoDict['DATE OF BIRTH'] = 'NA'
        else:
            dobReg = re.search(
                r'(\d{1,2}[-,./ ]*\s*((\d{1,2})|([a-zA-z]*))[-,./ ]\s*\d{4})|(\d{1,2}((st)|(nd)|(th)|(rd))[-,./ ]*\s*[a-zA-z]*[-,./ ]*\s*\d{4})|(d{1,2}[-,./ ]\s*\d{1,2}[-,./ ]\s*\d{4}|[a-zA-Z]*[-,./ ]\d{1,2}[-,./ ]\s*\d{4})',
                reg[0], re.IGNORECASE)
            dob = dobReg[0]
            infoDict['DATE OF BIRTH'] = dob.replace(',', ' ')

    #Function to get Gender.
    def getGender(self, inputString, infoDict, file):
        reg = re.search(r'((((gender)|(sex))[\s]*:\s*((male)|(female)|(m)|(f)))|((male)|(female)))', inputString, re.IGNORECASE)
        ext = file[-4:]
        if reg is None and ext == '.pdf':
            text = convertPDFToTextUsingPypdf2(file)  # checking for text using PyPDF2 if not found using pdfminer
            text = text.replace('\n', '')
            reg = re.search(r'((((gender)|(sex))[\s]*[-=:]\s*((male)|(female)|(m)|(f)))|((male)|(female)))',text, re.IGNORECASE)
            if reg is None:
                infoDict['GENDER'] = 'NA'
            else:
                g = re.search(r'((male)|(female)|(m)|(f))', reg[0], re.IGNORECASE)
                infoDict['GENDER'] = g[0].strip()
        elif reg is None:
            infoDict['GENDER'] = 'NA'
        else:
            g = re.search(r'((male)|(female)|(m)|(f))', reg[0], re.IGNORECASE)
            infoDict['GENDER'] = g[0].strip()

    #Function to get Nationality.
    def getNationality(self, inputString, infoDict, file):
        reg = re.search(r'((nationality)|(country))[\s]*[-=:]\s*[a-zA-Z]*', inputString, re.IGNORECASE)
        ext = file[-4:]
        if reg is None and ext == '.pdf':
            text = convertPDFToTextUsingPypdf2(file)  # checking for text using PyPDF2 if not found using pdfminer
            text = text.replace('\n', '')
            reg = re.search(r'(((nationality)|(country))[\s]*[-=:]\s*[a-zA-Z]*)',text, re.IGNORECASE)
            if reg is None:
                infoDict['NATIONALITY'] = 'NA'
            else:
                nl = re.search(r'([-=:]\s*[a-zA-Z]*)', reg[0], re.IGNORECASE)
                nl = re.split('[-=:]', nl[0])
                infoDict['NATIONALITY'] = nl[1].strip()
        elif reg is None:
            infoDict['NATIONALITY'] = 'NA'
        else:
            nl = re.search(r'([-=:]\s*[a-zA-Z]*)', reg[0], re.IGNORECASE)
            nl = re.split('[-=:]', nl[0])
            infoDict['NATIONALITY'] = nl[1].strip()

    #Function to get Current Address
    def getCurrentAddress(self, inputString, infoDict, file):
        reg = re.search(r'(((current location)|(current address))\s*[-:=]\s*[A-Za-z]*)|(place\s*[-=:]\s{1,5}[a-zA-z]+)', inputString,re.IGNORECASE)
        ext = file[-4:]
        if reg is None and ext == '.pdf':
            text = convertPDFToTextUsingPypdf2(file) # checking for text using PyPDF2 if not found using pdfminer
            text = text.replace('\n', '')
            reg = re.search(r'(((current location)|(current address))\s*[-:=]\s*[A-Za-z]*)|(place\s*[-=: ][ ]{1,5}[a-zA-z]+)', text, re.IGNORECASE)
            if reg is None:
                infoDict['CURRENT ADDRESS'] = 'NA'
            else:
                adr = re.search(r'([-=:]\s{0,5}[a-zA-z]+)', reg[0], re.IGNORECASE)
                adr = re.split('[-=:]', adr[0])
                infoDict['CURRENT ADDRESS'] = adr[1].strip()
        elif reg is None:
            infoDict['CURRENT ADDRESS'] = 'NA'
        else:
            adr = re.search(r'([-=:]\s{0,5}[a-zA-z]+)', reg[0], re.IGNORECASE)
            adr = re.split('[-=:]', adr[0])
            if adr[1].strip() == '':
                infoDict['CURRENT ADDRESS'] = 'NA'
            else:
                infoDict['CURRENT ADDRESS'] = adr[1].strip()

if __name__ == "__main__":
    p = Parse()