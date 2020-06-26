import os, glob, re
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import StringIO
from docx import Document
import PyPDF2


def convertDocxToText(path):
    txt = []
    document = Document(path)
    txt.append("\n".join([para.text for para in document.paragraphs]))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                txt.append("\n".join([para.text for para in cell.paragraphs]))

    return ('\n'.join([str(e) for e in txt]))



def convertPDFToText(path):
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, laparams=laparams)
    fp = open(path, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password = ""
    maxpages = 0
    caching = True
    pagenos = set()
    for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password, caching=caching,
                                  check_extractable=True):
        interpreter.process_page(page)
    fp.close()
    device.close()
    string = retstr.getvalue()
    retstr.close()
    string = string.replace('\n','')
    return string

def convertPDFToTextUsingPypdf2(path):
    pdfFileObj = open(path, 'rb')
    pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
    num_pages = pdfReader.numPages
    count = 0
    text = ""

    while count < num_pages:
        pageObj = pdfReader.getPage(count)
        count += 1
        text += pageObj.extractText()
    return text


def getData(inputString,file_path):
    dob1 = []
    gen = []
    nation = []
    loc = []

    reg = re.search(
        r'(((DOB)|(D.O.B)|(date of birth)|(birthday))\s*:*[- ]\s*\d{1,2}[-,./ ]*\s*((\d{1,2})|([a-zA-z]*))[-,./ ]\s*\d{4})|((DOB)|(D.O.B)|(date of birth)|(birthday))\s*:*-*\s*(\d{1,2}((st)|(nd)|(th)|(rd))[-,./ ]*\s*[a-zA-z]*[-,./ ]*\s*\d{4})|((DOB)|(D.O.B)|(Date of Birth)|(birthday))\s*:*-*\s*\d{1,2}[-,./ ]\s*\d{1,2}[-,./ ]\s*\d{4}|((DOB)|(D.O.B)|(Date of Birth)|(birthday))\s*:*-*\s*[a-zA-Z]*[-,./ ]\d{1,2}[-,./ ]\s*\d{4}',
        inputString, re.IGNORECASE)

    ext = file_path[-4:]
    if reg is None and ext == '.pdf':
        text = convertPDFToTextUsingPypdf2(file_path)  # checking for text using PyPDF2 if not found using pdfminer
        text = text.replace('\n', '')
        reg = re.search(
            r'(((DOB)|(D.O.B)|(date of birth)|(birthday))\s*:*[- ]\s*\d{1,2}[-,./ ]*\s*((\d{1,2})|([a-zA-z]*))[-,./ ]\s*\d{4})|((DOB)|(D.O.B)|(date of birth)|(birthday))\s*:*-*\s*(\d{1,2}((st)|(nd)|(th)|(rd))[-,./ ]*\s*[a-zA-z]*[-,./ ]*\s*\d{4})|((DOB)|(D.O.B)|(Date of Birth)|(birthday))\s*:*-*\s*\d{1,2}[-,./ ]\s*\d{1,2}[-,./ ]\s*\d{4}|((DOB)|(D.O.B)|(Date of Birth)|(birthday))\s*:*-*\s*[a-zA-Z]*[-,./ ]\d{1,2}[-,./ ]\s*\d{4}',
            text, re.IGNORECASE)

        if reg is None:
            dob1 += ['NA']
        else:
            dobReg = re.search(
                r'(\d{1,2}[-,./ ]*\s*((\d{1,2})|([a-zA-z]*))[-,./ ]\s*\d{4})|(\d{1,2}((st)|(nd)|(th)|(rd))[-,./ ]*\s*[a-zA-z]*[-,./ ]*\s*\d{4})|(d{1,2}[-,./ ]\s*\d{1,2}[-,./ ]\s*\d{4}|[a-zA-Z]*[-,./ ]\d{1,2}[-,./ ]\s*\d{4})',
                reg[0], re.IGNORECASE)
            dob1 += [dobReg[0].replace(',', ' ')]

    elif reg is None:
        dob1 += ['NA']
    else:
        dobReg = re.search(
            r'(\d{1,2}[-,./ ]*\s*((\d{1,2})|([a-zA-z]*))[-,./ ]\s*\d{4})|(\d{1,2}((st)|(nd)|(th)|(rd))[-,./ ]*\s*[a-zA-z]*[-,./ ]*\s*\d{4})|(d{1,2}[-,./ ]\s*\d{1,2}[-,./ ]\s*\d{4}|[a-zA-Z]*[-,./ ]\d{1,2}[-,./ ]\s*\d{4})',
            reg[0], re.IGNORECASE)
        dob = dobReg[0]
        dob1 += [dob.replace(',', ' ')]

# Extract Gender
    reg = re.search(r'((((gender)|(sex))[\s]*:\s*((male)|(female)|(m)|(f)))|((male)|(female)))', inputString, re.IGNORECASE)
    ext = file_path[-4:]
    if reg is None and ext == '.pdf':
        text = convertPDFToTextUsingPypdf2(file_path)  # checking for text using PyPDF2 if not found using pdfminer
        text = text.replace('\n', '')
        reg = re.search(r'((((gender)|(sex))[\s]*[-=:]\s*((male)|(female)|(m)|(f)))|((male)|(female)))',text, re.IGNORECASE)
        if reg is None:
            gen += ['NA']
        else:
            g = re.search(r'((male)|(female)|(m)|(f))', reg[0], re.IGNORECASE)
            gen += [g[0].strip()]
    elif reg is None:
        gen += ['NA']
    else:
        g = re.search(r'((male)|(female)|(m)|(f))', reg[0], re.IGNORECASE)
        gen += [g[0].strip()]


#Extract Nationality
    reg = re.search(r'((nationality)|(country))[\s]*[-=:]\s*[a-zA-Z]*', inputString, re.IGNORECASE)
    ext = file_path[-4:]
    if reg is None and ext == '.pdf':
        text = convertPDFToTextUsingPypdf2(file_path)  # checking for text using PyPDF2 if not found using pdfminer
        text = text.replace('\n', '')
        reg = re.search(r'(((nationality)|(country))[\s]*[-=:]\s*[a-zA-Z]*)',text, re.IGNORECASE)
        if reg is None:
            nation += ['NA']
        else:
            nl = re.search(r'([-=:]\s*[a-zA-Z]*)', reg[0], re.IGNORECASE)
            nl = re.split('[-=:]', nl[0])
            nation += [nl[1].strip()]
    elif reg is None:
        nation += ['NA']
    else:
        nl = re.search(r'([-=:]\s*[a-zA-Z]*)', reg[0], re.IGNORECASE)
        nl = re.split('[-=:]', nl[0])
        nation += [nl[1].strip()]

#Extract Current Address
    reg = re.search(r'(((current location)|(current address))\s*[-:=]\s*[A-Za-z]*)|(place\s*[-=:]\s{1,5}[a-zA-z]+)', inputString,re.IGNORECASE)
    ext = file_path[-4:]
    if reg is None and ext == '.pdf':
        text = convertPDFToTextUsingPypdf2(file_path) # checking for text using PyPDF2 if not found using pdfminer
        text = text.replace('\n', '')
        reg = re.search(r'(((current location)|(current address))\s*[-:=]\s*[A-Za-z]*)|(place\s*[-=: ][ ]{1,5}[a-zA-z]+)', text, re.IGNORECASE)
        if reg is None:
            loc += ['NA']
        else:
            adr = re.search(r'([-=:]\s{0,5}[a-zA-z]+)', reg[0], re.IGNORECASE)
            adr = re.split('[-=:]', adr[0])
            loc += [adr[1].strip()]
    elif reg is None:
        loc += ['NA']
    else:
        adr = re.search(r'([-=:]\s{0,5}[a-zA-z]+)', reg[0], re.IGNORECASE)
        adr = re.split('[-=:]', adr[0])
        if adr[1].strip() == '':
            loc += ['NA']
        else:
            loc += [adr[1].strip()]

    return dob1, gen, nation, loc

def readFile(fileName):

    extension = fileName.split(".")[-1]
    if extension == "pdf":
        return convertPDFToText(fileName)
    elif extension == "docx":
        try:
            return convertDocxToText(fileName)
        except:
            return ''
            pass
    else:
        print('Unsupported format')
        return ''
if __name__ == "__main__":
    # Glob module matches certain patterns
    doc_files = glob.glob("resumes/*.doc")
    docx_files = glob.glob("resumes/*.docx")
    pdf_files = glob.glob("resumes/*.pdf")

    files = set(doc_files + docx_files + pdf_files)
    files = list(files)
    print("%d files identified" % len(files))

    for file in files:
        print("Reading File %s" % file)
        inputString = readFile(file)
        dateOfbirth, gender, nationality, location = getData(inputString, file)
        print(dateOfbirth, gender, nationality, location)
